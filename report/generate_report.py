"""生成 MyAgent 项目学习报告 Word 文档"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ============================================================
# 全局样式设置
# ============================================================
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(4)

# 设置标题样式
for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = '微软雅黑'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if i == 1:
        hs.font.size = Pt(18)
        hs.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    elif i == 2:
        hs.font.size = Pt(15)
        hs.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    else:
        hs.font.size = Pt(13)
        hs.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)

# 页边距
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)


def add_image_placeholder(doc, caption: str, width: Inches = Inches(5)):
    """在文档中插入图片占位符（带边框的矩形 + 说明文字）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"【📷 请在此处粘贴图片：{caption}】")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.italic = True
    doc.add_paragraph()  # 空白行


def add_code_block(doc, code_text: str, language: str = "python"):
    """添加代码块（灰色底纹）"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    # 设置底纹
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F5F5F5')
    shd.set(qn('w:val'), 'clear')
    pPr.append(shd)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def add_bullet(doc, text: str, level: int = 0):
    """添加项目符号列表项"""
    p = doc.add_paragraph(text, style='List Bullet')
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.5 * (level + 1))
    return p


# ============================================================
# 封面
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("基于 Python 的 LLM Agent\n设计与实现")
run.font.size = Pt(28)
run.font.bold = True
run.font.name = '微软雅黑'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("—— MyAgent 项目学习报告")
run.font.size = Pt(16)
run.font.name = '微软雅黑'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
doc.add_paragraph()

# 封面信息表格
info_items = [
    ("课程名称", "人工智能基础"),
    ("项目名称", "MyAgent——LLM Agent 学习项目"),
    ("姓    名", "_______________"),
    ("学    号", "_______________"),
    ("专    业", "_______________"),
    ("日    期", "2026年6月1日"),
]
for label, value in info_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{label}：{value}")
    run.font.size = Pt(14)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_page_break()

# ============================================================
# 目录占位
# ============================================================
doc.add_heading('目  录', level=1)
add_image_placeholder(doc, "此处粘贴自动生成的目录，或在 Word 中「引用→目录→自动目录」生成")
doc.add_page_break()

# ============================================================
# 第一章 项目概述
# ============================================================
doc.add_heading('第一章  项目概述', level=1)

doc.add_heading('1.1 项目背景', level=2)
doc.add_paragraph(
    '随着大语言模型（LLM，Large Language Model）技术的飞速发展，AI Agent（人工智能代理）已成为人工智能领域的重要研究方向。'
    '传统的 LLM 聊天机器人只能进行文本对话，而 AI Agent 则能够调用外部工具（如搜索引擎、计算器、文件系统等），'
    '自主完成复杂任务。本项目旨在从零开始，使用 Python 语言构建一个具备工具调用、记忆管理、多会话管理等能力的 LLM Agent，'
    '深入理解 AI Agent 的核心原理与实现方式。'
)

doc.add_heading('1.2 项目目标', level=2)
add_bullet(doc, '掌握 OpenAI SDK 的使用方法，实现与 DeepSeek 大语言模型的交互')
add_bullet(doc, '实现 Agent 工具调用机制（Function Calling），使模型能够调用外部工具')
add_bullet(doc, '构建跨会话记忆系统，使 Agent 能够持久化用户偏好和重要信息')
add_bullet(doc, '实现多会话管理，支持会话存档、切换、删除')
add_bullet(doc, '实现对话压缩功能，在长对话中自动摘要旧消息以节省 Token 消耗')
add_bullet(doc, '通过 V1→V6 的迭代开发，记录完整的 Agent 演进过程')

doc.add_heading('1.3 技术栈', level=2)
# 技术栈表格
table = doc.add_table(rows=8, cols=3, style='Light Shading Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['技术/工具', '版本/说明', '用途']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.bold = True

data = [
    ('Python', '3.13', '主编程语言'),
    ('OpenAI SDK', '>=1.0.0', '调用 DeepSeek API（兼容 OpenAI 格式）'),
    ('DeepSeek V4', 'deepseek-v4-pro', '大语言模型（推理模式）'),
    ('python-dotenv', '>=1.0.0', '.env 环境变量管理，保护 API 密钥'),
    ('PyYAML', '>=6.0', '解析 config.yaml 配置文件'),
    ('BeautifulSoup4', '>=4.0', '网页内容解析'),
    ('Colorama', '>=0.4', '终端彩色输出'),
]
for row_idx, (tech, ver, purpose) in enumerate(data, start=1):
    for col_idx, val in enumerate((tech, ver, purpose)):
        table.rows[row_idx].cells[col_idx].text = val

doc.add_paragraph()
doc.add_heading('1.4 项目结构', level=2)
doc.add_paragraph('项目采用模块化设计，各组件职责清晰：')

add_code_block(doc, """MyAgent/
├── v6_agent.py                # 主程序（当前版本，~590 行）
├── config.yaml                # 配置文件（模型参数、工具开关、压缩设置）
├── requirements.txt           # Python 依赖包列表
├── .env                       # API 密钥（不提交到 Git）
├── README.md                  # 项目说明文档
├── tools/                     # 工具包（11 个工具）
│   ├── __init__.py            # 工具注册中心
│   ├── time_tool.py           # 时间查询
│   ├── weather_tool.py        # 天气查询（wttr.in）
│   ├── file_reader.py         # 文件读取
│   ├── file_writer.py         # 文件写入（沙箱模式）
│   ├── web_search.py          # 网页搜索（Bing）
│   ├── web_reader.py          # 网页全文阅读
│   ├── shell_tool.py          # 命令行执行（需确认）
│   ├── python_runner.py       # Python 代码执行（沙箱）
│   └── memory_tool.py         # 记忆管理（增删查）
├── examples/                  # V1~V5 版本示例 + 学习笔记
│   ├── NOTES.md               # 完整学习笔记（本文档重要参考）
│   ├── v1_single_chat.py      # V1：单次对话
│   ├── v2_loop_chat.py        # V2：循环多轮对话
│   ├── v3_with_tools.py       # V3：工具调用
│   ├── v4_agent.py            # V4：配置化 + 工具生态
│   └── v5_agent.py            # V5：记忆系统
└── workspace/                 # 运行时工作目录（不提交）
    ├── conversations/         # 会话存档（独立 JSON 文件）
    ├── memory.json            # 跨会话记忆持久化
    └── (Agent 创建的文件)
""")

add_image_placeholder(doc, "项目文件夹结构截图")

doc.add_page_break()

# ============================================================
# 第二章 版本迭代（V1→V6）
# ============================================================
doc.add_heading('第二章  版本迭代：从 V1 到 V6', level=1)
doc.add_paragraph(
    '本项目采用迭代开发模式，每个版本在前一版本的基础上增加新功能，逐步构建完整的 Agent 系统。'
    '下面详细描述每个版本的目标、核心改动、关键代码片段和学到的知识。'
)

# --- V1 ---
doc.add_heading('2.1  V1：单次对话（~25 行）', level=2)
doc.add_paragraph('目标：跑通 DeepSeek API，理解最基本的"发消息→收回复"流程。')

doc.add_heading('核心代码', level=3)
add_code_block(doc, """# V1 核心逻辑（简化）
from openai import OpenAI
from dotenv import load_dotenv
import os

load_dotenv()
client = OpenAI(
    api_key=os.getenv("DEEPSEEK_API_KEY"),
    base_url=os.getenv("DEEPSEEK_BASE_URL"),
)

question = input("IN [1]: ")
response = client.chat.completions.create(
    model=os.getenv("DEEPSEEK_MODEL"),
    messages=[{"role": "user", "content": question}]
)
print(f"OUT[1]: {response.choices[0].message.content}")""")

doc.add_heading('学到的知识', level=3)
add_bullet(doc, 'OpenAI SDK 的 client.chat.completions.create() 是真正发送 HTTP 请求的方法')
add_bullet(doc, 'messages 列表里每条消息有 role（user/assistant/system）和 content 两个字段')
add_bullet(doc, 'response.choices[0].message.content 用于提取模型的回复文本')
add_bullet(doc, '.env 文件 + python-dotenv 是最佳的密钥管理方案')
add_bullet(doc, 'DeepSeek API 完全兼容 OpenAI SDK 接口格式，切换成本极低')

add_image_placeholder(doc, "V1 运行截图：单次问答示例")

# --- V2 ---
doc.add_heading('2.2  V2：循环多轮对话（~90 行）', level=2)
doc.add_paragraph('目标：像网页版 ChatGPT 一样持续聊天，能记住上下文。')

doc.add_heading('V1→V2 关键改动', level=3)

table = doc.add_table(rows=6, cols=2, style='Light Shading Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.rows[0].cells[0].text = '改动项'
table.rows[0].cells[1].text = '说明'
for r in table.rows[0].cells:
    for p in r.paragraphs:
        for run in p.runs:
            run.font.bold = True

changes = [
    ('while True 循环', '替代一次性执行，实现持续对话'),
    ('messages 列表记忆', '将全部历史消息存在列表中，每轮全量发送给模型——这是"记忆"的最简单形态'),
    ('流式输出 (stream=True)', '实时显示模型的思考和回答过程，改善等待体验'),
    ('系统命令', '/exit、/clear、/help、/stats 四类命令，扩展了终端交互能力'),
    ('Token 估算统计', '中文字符 ×0.6 + 英文 ×0.3 的粗略估算，帮助了解用量'),
]
for row_idx, (item, desc) in enumerate(changes, start=1):
    table.rows[row_idx].cells[0].text = item
    table.rows[row_idx].cells[1].text = desc

doc.add_paragraph()
doc.add_heading('核心概念：对话记忆 = 一个不断增长的列表', level=3)
add_code_block(doc, """# V2 的记忆实现（最简形态）
messages = [{"role": "system", "content": "你是一个友好的AI助手"}]

while True:
    user_input = input("...")
    messages.append({"role": "user", "content": user_input})   # 追加用户消息

    response = client.chat.completions.create(
        model=model,
        messages=messages,   # 全量发送历史消息
        stream=True,
    )
    # ... 收集回答 ...
    messages.append({"role": "assistant", "content": answer})  # 追加模型回复""")

doc.add_heading('学到的知识', level=3)
add_bullet(doc, 'DeepSeek 支持前缀缓存（Prefix Caching），重复的历史部分只需约 1/10 价格')
add_bullet(doc, 'f-string（f"...{var}..."）比加号拼接更安全、更易读')
add_bullet(doc, '流式输出大幅改善用户体验，思考过程实时可见')

add_image_placeholder(doc, "V2 运行截图：多轮对话 + 思考过程展示")

# --- V3 ---
doc.add_heading('2.3  V3：工具调用（~140 行）', level=2)
doc.add_paragraph('目标：Agent 不仅能聊天，还能"做事"——调用计算器、查询时间。这是 Agent 区别于普通聊天机器人的核心能力。')

doc.add_heading('V2→V3 关键改动', level=3)
add_bullet(doc, '创建 tools/ 包：每个工具一个 Python 文件，包含实现函数和 JSON Schema 描述')
add_bullet(doc, 'tools/__init__.py 注册中心：统一管理 TOOL_FUNCTIONS 字典和 TOOL_DESCRIPTIONS 列表')
add_bullet(doc, 'API 调用加入 tools=TOOL_DESCRIPTIONS 参数，告知模型有哪些工具可用')
add_bullet(doc, '检查 msg.tool_calls 字段：若不为空则执行对应工具函数')
add_bullet(doc, '工具调用和结果用 role: "tool" + tool_call_id 加入消息历史')
add_bullet(doc, 'while True 循环处理多次连续工具调用（模型可能一轮调用多个工具或多轮调用）')

doc.add_heading('工具调用的完整流程', level=3)
add_code_block(doc, """# V3 工具调用循环（简化）
while True:
    response = client.chat.completions.create(
        model=model,
        messages=messages,
        tools=TOOL_DESCRIPTIONS,  # 关键：告诉模型有哪些工具
    )
    msg = response.choices[0].message

    if not msg.tool_calls:   # 不需要工具 → 退出循环，直接回答
        final_msg = msg
        break

    # 执行每个工具调用
    for tool_call in msg.tool_calls:
        tool_func = TOOL_FUNCTIONS[tool_call.function.name]
        tool_result = tool_func(**json.loads(tool_call.function.arguments))

    # 将工具调用和结果加入历史
    messages.append({"role": "assistant", "tool_calls": msg.tool_calls, ...})
    messages.append({"role": "tool", "tool_call_id": ..., "content": tool_result})""")

doc.add_heading('踩过的坑', level=3)
add_bullet(doc, 'Python 的 ^ 运算符是异或（XOR），不是幂运算——数学表达式中需转换为 **')
add_bullet(doc, 'DeepSeek 推理模式下，reasoning_content 必须原样传回，否则下次 API 调用报 400 错误（KV-cache 无法命中）')
add_bullet(doc, '模型可能连续调用多次工具，必须用 while 循环处理，不能只处理一次')

add_image_placeholder(doc, "V3 运行截图：工具调用过程（计算器 + 时间查询）")

doc.add_page_break()

# --- V4 ---
doc.add_heading('2.4  V4：配置化 + 工具生态（~180 行）', level=2)
doc.add_paragraph('目标：代码与配置分离，Agent 拥有丰富的工具生态——读写文件、联网搜索、执行代码。')

doc.add_heading('V3→V4 关键改动', level=3)

table = doc.add_table(rows=7, cols=2, style='Light Shading Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.rows[0].cells[0].text = '新增功能'
table.rows[0].cells[1].text = '实现方式'
for r in table.rows[0].cells:
    for p in r.paragraphs:
        for run in p.runs:
            run.font.bold = True

data_v4 = [
    ('config.yaml 配置外置', 'system prompt、模型参数、工具开关全部从 YAML 文件读取，改行为不需改代码'),
    ('文件沙箱', 'write_file 限定在 workspace_dir + allowed_paths 范围内；read_file 全开放'),
    ('命令行工具 (run_command)', '展示命令 + 说明，用户输入 confirm 确认后才执行，防误触'),
    ('联网搜索 (web_search)', '用 Bing 抓取搜索结果（DuckDuckGo 国内不可用）'),
    ('网页全文阅读 (read_webpage)', 'BeautifulSoup 解析网页正文'),
    ('Python 代码执行 (run_python)', '子进程沙箱执行，15 秒超时，替代了旧版计算器'),
]
for row_idx, (item, desc) in enumerate(data_v4, start=1):
    table.rows[row_idx].cells[0].text = item
    table.rows[row_idx].cells[1].text = desc

doc.add_paragraph()
doc.add_heading('设计亮点：配置驱动的工具注册', level=3)
doc.add_paragraph('所有工具统一在 tools/__init__.py 中注册，通过 config.yaml 的 tools 段控制开关。新增工具只需三步：')
add_bullet(doc, '创建新的 tool_xxx.py，实现功能函数 + JSON Schema 描述')
add_bullet(doc, '在 __init__.py 的 ALL_TOOLS 字典中注册')
add_bullet(doc, '在 config.yaml 中添加对应的开关项')

add_code_block(doc, """# tools/__init__.py 的设计
ALL_TOOLS = {
    "get_current_time": (time_tool.get_current_time, time_tool.DESCRIPTION),
    "get_weather":      (weather_tool.get_weather,   weather_tool.DESCRIPTION),
    # ... 更多工具 ...
}

def load_tools(config_path="config.yaml"):
    config = yaml.safe_load(open(config_path))
    tool_settings = config.get("tools", {})
    functions, descriptions = {}, []
    for name, (func, desc) in ALL_TOOLS.items():
        if tool_settings.get(name, False):  # YAML 中为 true 才启用
            functions[name] = func
            descriptions.append(desc)
    return functions, descriptions""")

doc.add_heading('踩过的坑', level=3)
add_bullet(doc, 'DuckDuckGo 搜索在国内网络环境下不可用，改用 Bing 抓取')
add_bullet(doc, 'API 的函数描述和本地函数都是一种"接口约定"——DESCRIPTION 就是给模型看的"菜单"')
add_bullet(doc, '推理模式下 temperature 参数无效，模型只认 reasoning_effort（low/medium/max）')
add_bullet(doc, '确认词用 confirm 而非简单的 y，防止误触执行危险命令')

add_image_placeholder(doc, "V4 运行截图：文件读写 + 联网搜索 + Python 执行")

# --- V5 ---
doc.add_heading('2.5  V5：记忆系统（~350 行）', level=2)
doc.add_paragraph(
    '目标：Agent 跨会话记住用户偏好和重要信息。此前 V2-V4 的"记忆"仅限于当前会话内的对话历史，'
    '退出程序后全部丢失。V5 引入持久化记忆，让 Agent 在下次启动时仍能"回忆起"之前的信息。'
)

doc.add_heading('V4→V5 关键改动', level=3)
add_bullet(doc, 'memory.json 持久化存储：以 JSON 键值对形式存储记忆，保存在 workspace 目录')
add_bullet(doc, '三个记忆工具：save_memory（保存）、list_memories（列出）、delete_memory（删除）')
add_bullet(doc, '启动时加载记忆，注入 system prompt："关于用户，你知道：..."——让模型自然知晓记忆内容')
add_bullet(doc, '手动记忆命令：/remember 主题: 内容、/forget 主题、/memories')
add_bullet(doc, '自动记忆：Agent 在对话中学到重要信息时，被 system prompt 引导主动调用 save_memory')
add_bullet(doc, '/clear 只清空对话历史，不删除记忆')

doc.add_heading('记忆 ≠ 对话历史', level=3)
doc.add_paragraph(
    '这是 V5 最重要的概念突破。对话历史是本次会话内的临时上下文，退出即消失；'
    '记忆是跨会话持久化的结构化信息，下次启动时自动加载。两者服务于不同目的：'
)
table = doc.add_table(rows=4, cols=3, style='Light Shading Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['特性', '对话历史', '跨会话记忆']):
    table.rows[0].cells[i].text = h
    for p in table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.font.bold = True
for row_idx, (k, v1, v2) in enumerate([
    ('生命周期', '当前会话内', '跨会话持久化'),
    ('存储位置', '内存中的 messages 列表', 'workspace/memory.json'),
    ('内容形式', '完整对话记录', '结构化键值对（主题→内容）'),
], start=1):
    table.rows[row_idx].cells[0].text = k
    table.rows[row_idx].cells[1].text = v1
    table.rows[row_idx].cells[2].text = v2

doc.add_paragraph()
doc.add_heading('重大技术挑战：XML 工具调用幻觉', level=3)
doc.add_paragraph(
    '这是 V5 开发中遇到的最棘手问题。模型有时不按 JSON Schema 规范调用工具，'
    '而是在文本回复中直接输出 XML 格式的 <function_calls><invoke name="..."> 标签。'
    '这是因为 DeepSeek 训练数据中混入了其他 Agent 框架的对话格式。'
)

doc.add_paragraph('解决方案：')
add_bullet(doc, '流式输出中用状态机实时拦截 XML 标签（开始→进入→结束），拦截部分不打印到屏幕')
add_bullet(doc, '输出完成后用正则表达式提取 XML 中的工具名和参数')
add_bullet(doc, '执行工具并将结果以标准 tool_call 格式加入消息历史')
add_bullet(doc, '通过 processing_tools 标志位重回工具处理循环，让模型读取工具结果')

add_image_placeholder(doc, "V5 运行截图：记忆保存 + 下次启动时加载记忆")

doc.add_page_break()

# --- V6 ---
doc.add_heading('2.6  V6：多会话管理 + 对话压缩（~590 行）', level=2)
doc.add_paragraph(
    '目标：每个会话独立存档，长对话自动压缩节省 Token。V6 是当前最新版本，'
    '在所有 V5 功能的基础上增加了完善的会话管理和智能压缩系统。'
)

doc.add_heading('V5→V6 核心改动', level=3)

doc.add_heading('（一）多会话管理', level=4)
add_bullet(doc, '独立存档：每个会话存为 workspace/conversations/<id>.json，不再将所有对话混在一起')
add_bullet(doc, 'sessions.json 清单：记录各会话的 ID、标题、轮数、更新时间')
add_bullet(doc, '会话命令：/new [标题]（新建）、/sessions（列表）、/load <ID>（加载）、/delete <ID>（删除）')
add_bullet(doc, '/save [标题] 手动保存，支持可选标题；/title <标题> 手动改名')
add_bullet(doc, '自动标题：首条用户消息的前 30 字符自动设为会话标题——低成本高收益的功能')
add_bullet(doc, '/clearall confirm 安全删除全部存档（必须输入 confirm 确认）')

doc.add_heading('（二）对话压缩', level=4)
doc.add_paragraph(
    '当用户与助手的对话轮数超过阈值（config.yaml 中 compression.max_messages，默认 30 条），'
    '保存时自动触发压缩：'
)
add_bullet(doc, '取最旧的用户+助手消息，调用 LLM 生成 500 字以内的中文摘要')
add_bullet(doc, '最近 12 条消息保留原文（确保最近的上下文精确可用）')
add_bullet(doc, '压缩后存档格式：{"summary": "摘要文本...", "messages": [最近消息]}')
add_bullet(doc, 'tool 消息必须随最近消息完整保留，否则加载后的对话无法继续工具调用')
add_bullet(doc, '加载时摘要注入 system prompt，并提示"请基于以上摘要继续对话，不要重复已完成的工作"')
add_bullet(doc, '兼容旧格式：旧版纯列表存档自动识别并正常加载')

doc.add_heading('压缩算法——倒序查找分界点', level=4)
add_code_block(doc, """# 从后往前数 keep_recent 条 user/assistant，找到精确的分界索引
count = 0
split_idx = len(messages)  # 默认全保留
for i in range(len(messages) - 1, 0, -1):
    if messages[i]["role"] in ("user", "assistant"):
        count += 1
        if count >= keep_recent:
            split_idx = i    # 分界点：从这里开始的所有消息保留原文
            break
# split_idx 之前的 user/assistant 消息 → 摘要
# split_idx 之后的所有消息（含 tool） → 原文保留""")

doc.add_heading('其他改进', level=4)
add_bullet(doc, '/history [N]：查看当前会话对话历史，可限定最近 N 轮')
add_bullet(doc, '修复 round_num 不递增的计数 Bug（XML 工具路径 continue 跳过 +1 的问题）')
add_bullet(doc, '.gitignore 增加 workspace/conversations/ 和 workspace/memory.json')

doc.add_heading('踩过的坑', level=3)
add_bullet(doc, '全量合并历史危害大：V5 启动时合并所有历史记录，对话越长越慢、越贵。改为独立存档 + 按需加载')
add_bullet(doc, '压缩时要保留 tool 消息：只压缩 user/assistant，tool 的调用和结果必须完整保留')
add_bullet(doc, '分界点须用倒序查找：从后往前数 keep_recent 条 user/assistant 来找分界索引，正序匹配内容容易因重复消息而出错')
add_bullet(doc, '老存档兼容：load_session_messages 返回 (summary, messages) 元组，但旧格式是纯列表，需用 isinstance(data, dict) 判断')
add_bullet(doc, '摘要注入 system prompt 时加"不要重复已完成的工作"至关重要，否则模型会重新执行历史中的任务')

add_image_placeholder(doc, "V6 运行截图：会话列表 (/sessions) + 对话压缩效果")
add_image_placeholder(doc, "V6 运行截图：/history 命令 + /load 加载会话")

doc.add_page_break()

# ============================================================
# 第三章 核心功能详解
# ============================================================
doc.add_heading('第三章  核心功能详解', level=1)

doc.add_heading('3.1 工具系统架构', level=2)
doc.add_paragraph(
    'MyAgent 拥有 11 个内置工具，涵盖时间查询、天气获取、文件读写、网页搜索、'
    '命令行执行、Python 沙箱和记忆管理。每个工具由两个部分组成：'
)
add_bullet(doc, '实现函数：接受参数、执行逻辑、返回字符串结果')
add_bullet(doc, 'JSON Schema 描述：定义函数名、参数类型和描述，供模型理解工具的用途')

doc.add_paragraph()
doc.add_paragraph('以下是一个完整的工具定义示例（天气查询）：')
add_code_block(doc, """# tools/weather_tool.py
import requests

DESCRIPTION = {
    "type": "function",
    "function": {
        "name": "get_weather",
        "description": "查询指定城市的天气信息",
        "parameters": {
            "type": "object",
            "properties": {
                "city": {"type": "string", "description": "城市名称，如 Beijing"}
            },
            "required": ["city"]
        }
    }
}

def get_weather(city: str) -> str:
    url = f"https://wttr.in/{city}?format=%C+%t+%h+%w"
    resp = requests.get(url, timeout=10)
    return resp.text if resp.status_code == 200 else "查询失败\"""")

doc.add_paragraph()
doc.add_heading('工具进化史', level=3)
table = doc.add_table(rows=5, cols=2, style='Light Shading Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['版本', '工具列表']):
    table.rows[0].cells[i].text = h
    for p in table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.font.bold = True
tools_evo = [
    ('V3（2个）', 'calculator, get_current_time'),
    ('V4（8个）', 'get_current_time, get_weather, read_file, write_file, web_search, read_webpage, run_command, run_python'),
    ('V5（11个）', 'V4 的 8 个 + save_memory, list_memories, delete_memory'),
    ('V6（11个）', '同 V5（工具无变化，增强会话管理 + 对话压缩）'),
]
for row_idx, (ver, tools) in enumerate(tools_evo, start=1):
    table.rows[row_idx].cells[0].text = ver
    table.rows[row_idx].cells[1].text = tools

doc.add_paragraph()

doc.add_heading('3.2 记忆系统原理', level=2)
doc.add_paragraph(
    '记忆系统是 Agent"个性化"的关键。每次启动时，从 memory.json 读取所有记忆，'
    '格式化为文本后拼接到 system prompt 中：'
)
add_code_block(doc, """# memory.json 示例
{
  "编程水平": "Python 学了1个月，会基础语法",
  "操作系统": "Windows 10",
  "偏好IDE": "VS Code"
}

# 注入 system prompt 后：
system_prompt = "你是一个有用的AI助手..."
system_prompt += "\\n\\n关于用户，你知道：\\n- 编程水平: Python 学了1个月\\n- 操作系统: Windows 10""")

doc.add_paragraph(
    'Agent 在对话中被 system prompt 引导，当识别到用户的重要偏好或背景信息时，'
    '会主动调用 save_memory 工具进行保存。用户也可以手动使用 /remember 命令保存记忆。'
)

doc.add_heading('3.3 会话管理设计', level=2)
doc.add_paragraph('V6 的会话管理采用"独立文件 + 清单索引"的模式：')
add_bullet(doc, '每个会话一个 JSON 文件（workspace/conversations/<id>.json），互不干扰')
add_bullet(doc, 'sessions.json 充当"目录"，记录所有会话的元信息（标题、轮数、时间）')
add_bullet(doc, '启动时不自动加载旧会话，用户根据需要 /load 指定会话')
add_bullet(doc, '保存时自动压缩长对话（超过阈值），压缩后的会话以 📦 标记')

add_image_placeholder(doc, "会话管理架构示意图（可画流程图：新建→保存→切换→加载→删除）")

doc.add_heading('3.4 对话压缩算法', level=2)
doc.add_paragraph('压缩流程如下：')
add_bullet(doc, '触发条件：保存会话时，统计 user+assistant 消息总数，若超过 max_messages（默认30条）则触发')
add_bullet(doc, '分界点计算：从消息列表末尾反向遍历，计数 user/assistant 消息，数够 keep_recent（默认12条）停止')
add_bullet(doc, '旧消息摘要：将分界点之前的 user/assistant 消息构建为纯文本对话，发送给 LLM 请求 500 字中文摘要')
add_bullet(doc, '新消息保留：分界点之后的所有消息（含 tool 消息）原样保留——确保工具调用链完整')
add_bullet(doc, '存档格式：{"summary": "...", "messages": [...]}，加载时摘要注入 system prompt')
doc.add_paragraph(
    '核心设计原则：压缩 user/assistant 对话内容，但绝不压缩 tool 消息。'
    '因为 tool 调用和结果是 API 对话链的必要组成部分，缺失会导致 loaded 对话无法继续。'
)

add_image_placeholder(doc, "对话压缩流程图（可画：触发→倒序查找分界点→LLM 摘要→新格式保存）")

doc.add_page_break()

# ============================================================
# 第四章 关键技术问题与解决方案
# ============================================================
doc.add_heading('第四章  关键技术问题与解决方案', level=1)
doc.add_paragraph('在 V1→V6 的开发过程中，遇到了许多实际工程问题。以下选取最具代表性的几个问题进行分析。')

doc.add_heading('4.1 XML 工具调用幻觉', level=2)
doc.add_paragraph('【问题描述】模型有时在文本回复中直接输出 <function_calls><invoke name="..."> 标签，而不是通过标准的 tool_calls 字段。')
doc.add_paragraph('【原因分析】DeepSeek 训练数据中混入了 Anthropic Claude 等 Agent 框架的对话格式，模型学到了"用 XML 标签调用工具"的模式。')
doc.add_paragraph('【解决方案】')
add_bullet(doc, '流式输出实时拦截：维护 in_xml 状态标志 + xml_buf 缓冲区，检测到 <function_calls> 或 <invoke 后切换状态，拦截后续内容不打印')
add_bullet(doc, '正则提取：用 _extract_xml_tool_calls() 函数在完整输出中匹配并提取工具名和参数')
add_bullet(doc, '标准化处理：将提取到的工具调用伪装成标准的 tool_calls 格式加入消息历史')
add_bullet(doc, '循环处理：通过 processing_tools 标志位重回工具处理循环，让模型读取执行结果')

doc.add_heading('4.2 reasoning_content 400 错误', level=2)
doc.add_paragraph('【问题描述】使用 DeepSeek 推理模式时，如果 assistant 消息缺少 reasoning_content，下次 API 调用会返回 HTTP 400 错误。')
doc.add_paragraph('【原因分析】DeepSeek 服务端按完整消息序列维护 KV-cache。如果某条 assistant 消息缺少了之前的 reasoning_content，缓存无法命中，服务端拒绝请求。')
doc.add_paragraph('【解决方案】每次保存 assistant 消息时，检查 msg.reasoning_content 是否存在，若存在则一并存入历史。')

doc.add_heading('4.3 round_num 计数 Bug', level=2)
doc.add_paragraph('【问题描述】V6 重构后发现 round_num 一直为 1，不会递增。')
doc.add_paragraph('【原因分析】XML 工具处理路径在执行完工具后使用 continue 跳过输入环节，但如果 +1 放在 continue 之后则永远不会执行。同时，正常完成路径也需要正确放置 +1。')
doc.add_paragraph('【解决方案】只在正常完成路径（无 XML 工具调用）的末尾执行 round_num += 1；XML 工具路径的 continue 自然跳过该语句。')

doc.add_heading('4.4 全量合并历史的性能问题', level=2)
doc.add_paragraph('【问题描述】V5 启动时将 session.json 中所有历史消息合并到 messages 列表中。随着对话积累，启动越来越慢，每次 API 调用也越来越贵。')
doc.add_paragraph('【解决方案（V6）】改为独立文件存档 + 按需加载：每个会话独立 JSON 文件，启动时从空白开始，用户用 /load 命令加载指定会话。同时引入对话压缩机制，长对话自动摘要旧消息。')

doc.add_page_break()

# ============================================================
# 第五章 项目展示
# ============================================================
doc.add_heading('第五章  项目成果展示', level=1)

doc.add_heading('5.1 系统命令一览', level=2)
table = doc.add_table(rows=14, cols=2, style='Light Shading Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['命令', '功能']):
    table.rows[0].cells[i].text = h
    for p in table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.font.bold = True
commands = [
    ('/help', '显示帮助信息'),
    ('/exit', '保存当前会话并退出'),
    ('/save [标题]', '手动保存当前会话（可选标题）'),
    ('/new [标题]', '保存并开始新会话'),
    ('/sessions', '查看所有已保存的会话'),
    ('/load <ID>', '加载指定会话'),
    ('/delete <ID>', '删除指定会话'),
    ('/title <标题>', '设置当前会话标题'),
    ('/clear', '清空当前会话对话'),
    ('/clearall confirm', '删除全部会话存档'),
    ('/history [N]', '查看对话历史（可选最近N轮）'),
    ('/remember 主题: 内容', '手动保存记忆'),
    ('/memories', '查看所有记忆'),
]
for row_idx, (cmd, desc) in enumerate(commands, start=1):
    table.rows[row_idx].cells[0].text = cmd
    table.rows[row_idx].cells[1].text = desc

doc.add_paragraph()

doc.add_heading('5.2 工具能力展示', level=2)
doc.add_paragraph('以下是 Agent 各工具的典型使用场景：')
add_bullet(doc, '🌤️ 天气查询："今天北京天气怎么样？" → get_weather("Beijing")')
add_bullet(doc, '📁 文件操作："帮我写一个 Hello World 的 Python 脚本" → write_file("hello.py", ...)')
add_bullet(doc, '🔍 联网搜索："搜索最新的人工智能新闻" → web_search("AI news") + read_webpage(url)')
add_bullet(doc, '💻 代码执行："帮我算一下斐波那契数列第30项" → run_python("...")')
add_bullet(doc, '🖥️ 命令行："帮我创建一个新文件夹" → run_command("mkdir new_folder")（需 confirm 确认）')
add_bullet(doc, '🧠 记忆："记住我喜欢用 VS Code" → save_memory("偏好IDE", "VS Code")')

doc.add_heading('5.3 运行效果展示', level=2)
doc.add_paragraph('（请在此处粘贴以下运行截图）')
add_bullet(doc, '截图 1：程序启动界面（会话ID、已有存档数量）', level=1)
add_bullet(doc, '截图 2：Agent 使用工具查询天气的全过程（思考→工具调用→结果→回答）', level=1)
add_bullet(doc, '截图 3：/sessions 命令的会话列表展示', level=1)
add_bullet(doc, '截图 4：/history 命令的对话历史展示', level=1)
add_bullet(doc, '截图 5：记忆保存与跨会话加载的对比（保存记忆→退出→重新启动→Agent 主动引用记忆）', level=1)
add_bullet(doc, '截图 6：长对话压缩效果（/sessions 中的 📦 标记 + 加载时显示摘要）', level=1)

doc.add_page_break()

# ============================================================
# 第六章 总结与展望
# ============================================================
doc.add_heading('第六章  总结与展望', level=1)

doc.add_heading('6.1 项目总结', level=2)
doc.add_paragraph(
    '通过 MyAgent 项目的六个版本迭代开发，我从零开始完整构建了一个具备实用价值的 LLM Agent 系统。'
    '主要收获包括：'
)

add_bullet(doc, '深入理解了 LLM Agent 的核心架构：对话记忆 → 工具调用 → 记忆系统 → 会话管理 → 对话压缩，每一步都建立在前一步的基础之上')
add_bullet(doc, '掌握了 OpenAI SDK 的使用方法：消息构建、流式输出、Function Calling、推理模式等核心 API')
add_bullet(doc, '学会了工程化的配置管理：YAML 配置外置、.env 密钥管理、模块化工具注册')
add_bullet(doc, '积累了 LLM 应用的实战经验：XML 幻觉处理、推理模式兼容、Token 优化、对话压缩等')
add_bullet(doc, '理解了"记忆"与"对话历史"的本质区别：前者是跨会话持久化的结构化知识，后者是当前会话的临时上下文')
add_bullet(doc, '实践了迭代开发方法：每个版本功能单一、代码量可控（从 25 行到 590 行），逐步演进而非一步到位')

doc.add_heading('6.2 技术亮点', level=2)
add_bullet(doc, '配置驱动的工具注册系统：新增工具只需三步，无需修改主程序代码')
add_bullet(doc, 'XML 工具调用的容错处理：状态机拦截 + 正则提取，保障 Agent 在非标准输出下仍能正常工作')
add_bullet(doc, '对话压缩的倒序查找算法：精确保留最近 N 条消息 + 完整保留 tool 调用链')
add_bullet(doc, 'system prompt 注入的记忆加载：无需额外 API 调用，模型"自然"地知晓用户信息')
add_bullet(doc, '安全设计：命令确认机制、文件沙箱、密钥保护——多个环节考虑了安全性')

doc.add_heading('6.3 不足与改进方向', level=2)
add_bullet(doc, 'Token 精确计数：目前使用中文字符 ×0.6 的粗略估算，未使用 tiktoken 等精确分词工具')
add_bullet(doc, '工具调用迭代限制：当前 while True 可能无限循环，需要添加最大迭代次数限制')
add_bullet(doc, '错误恢复机制：工具执行失败时缺少自动重试或降级策略')
add_bullet(doc, '多模态支持：目前仅支持文本，未来可扩展图片理解、语音交互')
add_bullet(doc, '计划-执行（Plan-then-Execute）模式：让 Agent 先制定计划再逐步执行，提高复杂任务的完成率')
add_bullet(doc, 'Web UI：目前仅支持终端交互，未来可开发 Web 界面')

doc.add_heading('6.4 学习心得', level=2)
doc.add_paragraph(
    '这个项目是我大一下学期"人工智能基础"课程的学习成果。在为期数周的开发过程中，我最大的体会是：'
    '"Agent 不是魔法，而是一层一层叠加的工程能力。" V1 只会说一句话，V2 能记住上下文，'
    'V3 能调用工具，V4 能读写文件和联网，V5 能跨会话记忆，V6 能管理多轮对话并智能压缩。'
    '每一步的代码量都不多（最多增加 ~200 行），但每一步都解决了一个实际问题。'
)
doc.add_paragraph(
    '另一个重要收获是：好的 AI 应用不仅仅是调用 API，还需要大量的"防御性设计"——'
    '处理 XML 幻觉、确保 reasoning_content 回传、compatibility 兼容旧格式、'
    '安全的命令确认机制……这些"脏活累活"才是让 Agent 稳定运行的关键。'
)
doc.add_paragraph(
    '感谢 DeepSeek 提供了强大且实惠的 API，让个人开发者也能用上顶级的推理模型。'
    '感谢 VS Code + GitHub Copilot 提供的 AI 辅助编程体验，让学习效率大幅提升。'
)

doc.add_paragraph()
doc.add_paragraph()

# ============================================================
# 参考文献
# ============================================================
doc.add_heading('参考文献', level=1)
refs = [
    '[1] OpenAI. OpenAI Python SDK Documentation. https://github.com/openai/openai-python',
    '[2] DeepSeek. DeepSeek API 文档. https://api-docs.deepseek.com/',
    '[3] Python-dotenv. Python-dotenv Documentation. https://github.com/theskumar/python-dotenv',
    '[4] PyYAML. PyYAML Documentation. https://pyyaml.org/',
    '[5] BeautifulSoup4. BeautifulSoup Documentation. https://www.crummy.com/software/BeautifulSoup/',
    '[6] 本项目 GitHub 仓库：https://github.com/LiuAn-HuaMing/MyAgent',
]
for ref in refs:
    doc.add_paragraph(ref)


# ============================================================
# 保存文档
# ============================================================
output_path = os.path.join(os.path.dirname(__file__), "MyAgent学习报告.docx")
doc.save(output_path)
print(f"✅ 学习报告已生成：{output_path}")
